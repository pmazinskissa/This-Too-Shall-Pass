from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


class DocxExporter:
    """Export meeting summaries to Word documents"""

    def export(self, summary: Dict[str, Any], output_path: str,
               title: str = "", date: str = "", duration: str = "") -> None:
        """
        Export a meeting summary to a Word document

        Args:
            summary: Dictionary containing structured summary sections
            output_path: Path to save the Word document
            title: Meeting title
            date: Meeting date
            duration: Meeting duration
        """
        # Create document
        doc = Document()

        # Configure document styles
        self._setup_document_styles(doc)

        # Add header
        self._add_header(doc, title, date, duration)

        # If we have the markdown, use that directly
        if 'markdown' in summary and summary['markdown']:
            self._add_from_markdown(doc, summary['markdown'])

            # Check if key quotes were properly added from markdown
            # If not, add them from structured data as fallback
            if summary.get("key_quotes") and self._should_add_quotes_fallback(summary['markdown']):
                print("Adding key quotes from structured data as fallback...")
                doc.add_heading("Key Quotes", level=1)
                self._add_key_quotes_structured(doc, summary.get("key_quotes", []))
        else:
            # Otherwise use the previously extracted structured data
            # Add executive summary
            self._add_executive_summary(doc, summary.get("executive_summary", ""))

            # Add participants table
            self._add_participants_table(doc, summary.get("participants", []))

            # Add detailed summary
            self._add_detailed_summary(doc, summary.get("detailed_summary", []))

            # Add decisions table
            self._add_decisions_table(doc, summary.get("decisions_made", []))

            # Add actions table
            self._add_actions_table(doc, summary.get("actions_planned", []))

            # Add open questions table
            self._add_open_questions_table(doc, summary.get("open_questions", []))

            # Add key quotes
            self._add_key_quotes(doc, summary.get("key_quotes", []))

            # Add sentiment analysis
            self._add_sentiment_analysis(doc, summary.get("sentiment_analysis", ""))

            # Add content gaps
            self._add_content_gaps(doc, summary.get("content_gaps", []))

            # Add technical terminology
            self._add_terminology(doc, summary.get("terminology", []))

        # Save document
        doc.save(output_path)

    def _add_key_quotes_from_markdown(self, doc: Document, section_content: str) -> None:
        """
        Extract and add key quotes from markdown content with multiple format support

        Args:
            doc: The Document object
            section_content: The markdown content for the Key Quotes section
        """
        import re

        # Debug: Print the section content to understand the format
        print(f"Key Quotes section content:\n{repr(section_content)}")

        quotes_added = False

        # Try multiple quote patterns to handle different formats
        quote_patterns = [
            # Standard blockquote format: > "quote" – attribution
            r'>\s*"([^"]+)"\s*–\s*(.+?)(?=\n>|\n\n|$)',
            # Blockquote with em dash: > "quote" — attribution
            r'>\s*"([^"]+)"\s*—\s*(.+?)(?=\n>|\n\n|$)',
            # Blockquote without quotes: > quote – attribution
            r'>\s*([^"–—]+?)\s*–\s*(.+?)(?=\n>|\n\n|$)',
            # Simple format: "quote" – attribution (no >)
            r'"([^"]+)"\s*–\s*(.+?)(?=\n|$)',
            # Simple format with em dash: "quote" — attribution
            r'"([^"]+)"\s*—\s*(.+?)(?=\n|$)',
        ]

        for pattern in quote_patterns:
            quotes = re.findall(pattern, section_content, re.MULTILINE | re.DOTALL)
            if quotes:
                print(f"Found {len(quotes)} quotes with pattern: {pattern}")
                for quote_text, attribution in quotes:
                    quote_text = quote_text.strip()
                    attribution = attribution.strip()

                    # Skip empty quotes
                    if quote_text and attribution:
                        p = doc.add_paragraph(style='Quote')
                        p.add_run(f'"{quote_text}"')
                        p.add_run(f" — {attribution}")
                        quotes_added = True
                break  # Use the first pattern that finds quotes

        # If no quotes found, try to extract any text that looks like quotes
        if not quotes_added:
            # Look for any quoted text in the section
            quote_lines = []
            for line in section_content.split('\n'):
                line = line.strip()
                # Skip empty lines and markdown artifacts
                if line and not line.startswith('#') and not re.match(r'^[\s\-|:>]+$', line):
                    quote_lines.append(line)

            if quote_lines:
                print(f"Adding {len(quote_lines)} quote lines as fallback")
                for line in quote_lines:
                    # Clean up the line
                    line = re.sub(r'^>\s*', '', line)  # Remove blockquote markers
                    if line.strip():
                        p = doc.add_paragraph(style='Quote')
                        p.add_run(line.strip())
                        quotes_added = True

        # If still no quotes, add a note
        if not quotes_added:
            print("No quotes found, adding fallback message")
            doc.add_paragraph("No notable quotes recorded.")

    def _should_add_quotes_fallback(self, markdown_text: str) -> bool:
        """
        Check if the Key Quotes section in markdown appears to be empty or malformed

        Args:
            markdown_text: The full markdown text

        Returns:
            True if we should use structured data fallback for quotes
        """
        import re

        # Find the Key Quotes section
        section_pattern = r'##\s+\d+\.\s+Key Quotes\s*\n(.*?)(?=##\s+\d+\.|$)'
        match = re.search(section_pattern, markdown_text, re.DOTALL)

        if not match:
            return True  # No Key Quotes section found

        quotes_content = match.group(1).strip()

        # Check if the section is essentially empty or only contains "No notable quotes"
        if not quotes_content or "No notable quotes" in quotes_content or len(quotes_content) < 20:
            return True

        # Check if there are actual quote markers (> or ")
        if '>' not in quotes_content and '"' not in quotes_content:
            return True

        return False

    def _add_key_quotes_structured(self, doc: Document, quotes: List[Any]) -> None:
        """
        Add key quotes from structured data (not markdown)

        Args:
            doc: The Document object
            quotes: List of quote dictionaries or strings
        """
        if not quotes:
            doc.add_paragraph("No notable quotes recorded.")
            return

        for quote in quotes:
            if isinstance(quote, dict) and "quote" in quote:
                p = doc.add_paragraph(style='Quote')
                p.add_run(f'"{quote["quote"]}"')
                if "attribution" in quote and quote["attribution"]:
                    p.add_run(f" — {quote['attribution']}")
            elif isinstance(quote, str):
                p = doc.add_paragraph(style='Quote')
                # If it's already quoted, use as-is, otherwise add quotes
                if quote.startswith('"') and quote.endswith('"'):
                    p.add_run(quote)
                else:
                    p.add_run(f'"{quote}"')

    def _add_from_markdown(self, doc: Document, markdown_text: str) -> None:
        """
        Parse markdown and add it to the document with appropriate formatting

        Args:
            doc: The Document object
            markdown_text: The markdown text to add
        """
        import re

        # Split the markdown by section headers
        section_pattern = r'##\s+\d+\.\s+(.*?)\n(.*?)(?=##\s+\d+\.|$)'
        section_matches = re.findall(section_pattern, markdown_text, re.DOTALL)

        for section_title, section_content in section_matches:
            section_title = section_title.strip()

            # Add section heading
            heading = doc.add_heading(section_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Clean the section content - remove table separators and other markdown artifacts
            section_content = self._clean_markdown_content(section_content)

            # Process based on section type
            if "Participants" in section_title or "Decisions Made" in section_title or \
                    "Actions Planned" in section_title or "Open Questions" in section_title or \
                    "Technical Terminology" in section_title:
                # This is a table section
                self._add_markdown_table(doc, section_content)

            elif "Conversation Flow" in section_title:
                # Extract and add conversation flow subsections
                scene_pattern = r'###\s+\d+\s+·\s+(.*?)\n(.*?)(?=###\s+\d+|$)'
                scenes = re.findall(scene_pattern, section_content, re.DOTALL)

                for title, content in scenes:
                    doc.add_heading(title.strip(), level=2)
                    # Clean the content before adding
                    clean_content = self._clean_markdown_content(content)
                    doc.add_paragraph(clean_content.strip())

            elif "Key Quotes" in section_title:
                # Handle multiple quote formats more robustly
                self._add_key_quotes_from_markdown(doc, section_content)

            elif "Content Gaps" in section_title:
                # Extract bullet points
                bullet_pattern = r'-\s*(.*?)(?=\n-|\n\n|$)'
                gaps = re.findall(bullet_pattern, section_content)

                for gap in gaps:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(gap.strip())

            else:
                # Regular paragraph section (Executive Summary, Sentiment Analysis)
                # Clean content and split by paragraphs to handle it properly
                paragraphs = [p for p in section_content.split('\n\n') if p.strip()]
                for para in paragraphs:
                    # Skip lines that look like table separators or other markdown formatting
                    if not re.match(r'^[\s\-|:]+$', para.strip()):
                        doc.add_paragraph(para.strip())

    def _clean_markdown_content(self, content: str) -> str:
        """
        Clean markdown content by removing table separators and other markdown artifacts

        Args:
            content: The markdown content to clean

        Returns:
            Cleaned markdown content
        """
        import re

        # Remove table separator lines (lines with only |, -, :, and spaces)
        lines = content.split('\n')
        cleaned_lines = []

        for line in lines:
            # Skip table separator lines
            if re.match(r'^\s*(\|[\s\-:]+\|)+\s*$', line):
                continue
            # Skip empty separator lines with just dashes
            if re.match(r'^\s*-{3,}\s*$', line):
                continue
            cleaned_lines.append(line)

        # Join the cleaned lines back together
        cleaned_content = '\n'.join(cleaned_lines)

        # Remove any trailing markdown artifact lines
        cleaned_content = re.sub(r'\n\s*[\-\|][\s\-\|:]*$', '', cleaned_content)

        return cleaned_content

    def _add_markdown_table(self, doc: Document, markdown_table_text: str) -> None:
        """
        Convert a markdown table to a Word table with improved parsing

        Args:
            doc: The Document object
            markdown_table_text: The markdown table text
        """
        import re

        # If the text doesn't contain any table markers, just add it as text
        if "|" not in markdown_table_text:
            doc.add_paragraph(markdown_table_text)
            return

        # Split the content by lines and identify table rows
        lines = [line for line in markdown_table_text.strip().split('\n') if line.strip()]

        # Need at least header row and separator to proceed
        if len(lines) < 2:
            doc.add_paragraph(markdown_table_text)
            return

        # Get header line to determine columns
        header_line = lines[0]

        # Count the number of columns by counting pipe characters and subtracting 1
        # (accounting for the beginning and ending pipe)
        num_columns = header_line.count('|') - 1

        # If we can't determine columns properly, fall back to text
        if num_columns <= 0:
            doc.add_paragraph(markdown_table_text)
            return

        # Create a new table with the determined number of columns
        table = doc.add_table(rows=1, cols=num_columns)
        table.style = 'Table Grid'

        # Process each line as a row, explicitly removing separator rows
        row_idx = 0
        for line in lines:
            # Skip separator rows (containing dashes)
            if re.match(r'\s*\|(\s*[-:]+\s*\|)+\s*$', line):
                continue

            # Clean the line and split by pipe character
            line = line.strip()
            if line.startswith('|'):
                line = line[1:]
            if line.endswith('|'):
                line = line[:-1]

            # Split into cells and ensure we have the right number
            cells = line.split('|')

            # Skip rows that don't have enough cells (might be leftover separator data)
            if len(cells) < num_columns:
                continue

            # Create a new row if this isn't the header
            if row_idx == 0:
                # Header row already exists
                header_row = table.rows[0]
            else:
                # Add a new row for data
                row_cells = table.add_row().cells

            # Add cell contents
            for col_idx, cell_content in enumerate(cells):
                if col_idx < num_columns:  # Ensure we don't go beyond our columns
                    cell_text = cell_content.strip()

                    # Skip cells that only contain dashes (part of separators)
                    if re.match(r'^-+$', cell_text):
                        continue

                    if row_idx == 0:
                        # Header cells
                        header_row.cells[col_idx].text = cell_text
                        # Make header bold
                        for paragraph in header_row.cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    else:
                        # Data cells
                        row_cells[col_idx].text = cell_text

            row_idx += 1

        # No need to add an empty paragraph after each table - REMOVED

    def _setup_document_styles(self, doc: Document) -> None:
        """Set up document styles"""
        # Default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        # Heading styles
        for i in range(1, 4):
            style = doc.styles[f'Heading {i}']
            font = style.font
            font.name = 'Times New Roman'
            font.bold = True
            if i == 1:
                font.size = Pt(16)
            elif i == 2:
                font.size = Pt(14)
            else:
                font.size = Pt(12)

        # Create Quote style if it doesn't exist
        if 'Quote' not in doc.styles:
            quote_style = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.base_style = doc.styles['Normal']
            quote_style.font.italic = True
            quote_paragraph_format = quote_style.paragraph_format
            quote_paragraph_format.left_indent = Inches(0.5)
            quote_paragraph_format.space_before = Pt(6)
            quote_paragraph_format.space_after = Pt(6)

    def _add_header(self, doc: Document, title: str, date: str, duration: str) -> None:
        """Add document header with meeting info"""
        # Create the title with "Summary" appended
        full_title = f"{title} Summary"

        # Add centered, underlined title
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading_para.add_run(full_title)
        heading_run.bold = True
        heading_run.underline = True
        heading_run.font.size = Pt(16)  # Adjust size as needed
        heading_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue - modify these values as needed

        # Add metadata in italics on a new line
        metadata = ""
        if date:
            metadata += f"{date}"
        if date and duration:
            metadata += f" • {duration}"
        elif duration:
            metadata += f"{duration}"

        if metadata:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(metadata)
            run.italic = True

        # Add some space after header
        doc.add_paragraph()

    def _add_executive_summary(self, doc: Document, executive_summary: str) -> None:
        """Add executive summary section"""
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(executive_summary)
        # Removed extra spacing

    def _add_participants_table(self, doc: Document, participants: List[Dict[str, str]]) -> None:
        """Add participants table"""
        doc.add_heading("Participants", level=1)

        if not participants:
            doc.add_paragraph("No participant information available.")
            return

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Name"
        header_cells[1].text = "Organization / Title"
        header_cells[2].text = "Meeting Role"

        # Make header row bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add participant rows
        for participant in participants:
            if isinstance(participant, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = participant.get("name", "")
                row_cells[1].text = participant.get("organization", "")
                row_cells[2].text = participant.get("role", "")
            else:
                # Handle string entries
                row_cells = table.add_row().cells
                row_cells[0].text = str(participant)
                row_cells[1].text = ""
                row_cells[2].text = ""

        # Removed extra spacing

    def _add_detailed_summary(self, doc: Document, detailed_summary: List[Dict[str, str]]) -> None:
        """Add detailed summary section with subsections"""
        doc.add_heading("Conversation Flow Summary", level=1)

        if not detailed_summary:
            doc.add_paragraph("No detailed summary available.")
            return

        # Handle different data structures
        if isinstance(detailed_summary, list):
            for section in detailed_summary:
                if isinstance(section, dict) and "title" in section and "content" in section:
                    doc.add_heading(section["title"], level=2)
                    doc.add_paragraph(section["content"])
                elif isinstance(section, str):
                    doc.add_paragraph(section)
        elif isinstance(detailed_summary, dict):
            for title, content in detailed_summary.items():
                doc.add_heading(title, level=2)
                doc.add_paragraph(content)
        elif isinstance(detailed_summary, str):
            doc.add_paragraph(detailed_summary)

        # Removed extra spacing

    def _add_decisions_table(self, doc: Document, decisions: List[Any]) -> None:
        """Add decisions table"""
        doc.add_heading("Decisions Made", level=1)

        if not decisions:
            doc.add_paragraph("No decisions recorded.")
            return

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Decision"
        header_cells[1].text = "Details"
        header_cells[2].text = "Owner(s)"

        # Make header row bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add decision rows
        for decision in decisions:
            if isinstance(decision, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = decision.get("decision", "")
                row_cells[1].text = decision.get("details", "")
                row_cells[2].text = decision.get("owner", "")
            elif isinstance(decision, str):
                row_cells = table.add_row().cells
                row_cells[0].text = decision
                row_cells[1].text = ""
                row_cells[2].text = ""

        # Removed extra spacing

    def _add_actions_table(self, doc: Document, actions: List[Any]) -> None:
        """Add actions table"""
        doc.add_heading("Actions Planned", level=1)

        if not actions:
            doc.add_paragraph("No actions recorded.")
            return

        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Action"
        header_cells[1].text = "Responsible"
        header_cells[2].text = "Timeline"
        header_cells[3].text = "Notes"

        # Make header row bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add action rows
        for action in actions:
            if isinstance(action, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = action.get("action", "")
                row_cells[1].text = action.get("responsible", "")
                row_cells[2].text = action.get("timeline", "")
                row_cells[3].text = action.get("notes", "")
            elif isinstance(action, str):
                row_cells = table.add_row().cells
                row_cells[0].text = action
                row_cells[1].text = ""
                row_cells[2].text = ""
                row_cells[3].text = ""

        # Removed extra spacing

    def _add_open_questions_table(self, doc: Document, questions: List[Any]) -> None:
        """Add open questions table"""
        doc.add_heading("Open Questions", level=1)

        if not questions:
            doc.add_paragraph("No open questions recorded.")
            return

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Question"
        header_cells[1].text = "Context"
        header_cells[2].text = "Owner"

        # Make header row bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add question rows
        for question in questions:
            if isinstance(question, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = question.get("question", "")
                row_cells[1].text = question.get("context", "")
                row_cells[2].text = question.get("owner", "")
            elif isinstance(question, str):
                row_cells = table.add_row().cells
                row_cells[0].text = question
                row_cells[1].text = ""
                row_cells[2].text = ""

        # Removed extra spacing

    def _add_key_quotes(self, doc: Document, quotes: List[Any]) -> None:
        """Add key quotes section (original method for structured data)"""
        doc.add_heading("Key Quotes", level=1)

        if not quotes:
            doc.add_paragraph("No notable quotes recorded.")
            return

        for quote in quotes:
            if isinstance(quote, dict) and "quote" in quote:
                p = doc.add_paragraph(style='Quote')
                p.add_run(f'"{quote["quote"]}"')
                if "attribution" in quote and quote["attribution"]:
                    p.add_run(f" — {quote['attribution']}")
            elif isinstance(quote, str):
                p = doc.add_paragraph(style='Quote')
                # If it's already quoted, use as-is, otherwise add quotes
                if quote.startswith('"') and quote.endswith('"'):
                    p.add_run(quote)
                else:
                    p.add_run(f'"{quote}"')

        # Removed extra spacing

    def _add_sentiment_analysis(self, doc: Document, sentiment: str) -> None:
        """Add sentiment analysis section"""
        doc.add_heading("Sentiment Analysis", level=1)
        doc.add_paragraph(sentiment)
        # Removed extra spacing

    def _add_content_gaps(self, doc: Document, gaps: List[str]) -> None:
        """Add potential content gaps section"""
        doc.add_heading("Content Gaps", level=1)

        if not gaps:
            doc.add_paragraph("No content gaps identified.")
            return

        # Add bulleted list
        for gap in gaps:
            if isinstance(gap, dict) and "gap" in gap:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{gap['gap']}")
            elif isinstance(gap, str):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{gap}")

        # Removed extra spacing

    def _add_terminology(self, doc: Document, terms: List[Any]) -> None:
        """Add technical terminology section"""
        doc.add_heading("Technical Terminology & Acronyms", level=1)

        if not terms:
            doc.add_paragraph("No technical terminology recorded.")
            return

        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Term"
        header_cells[1].text = "Definition"

        # Make header row bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add term rows
        for term in terms:
            if isinstance(term, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = term.get("term", "")
                row_cells[1].text = term.get("definition", "")
            elif isinstance(term, str) and ":" in term:
                # Handle "term: definition" format
                parts = term.split(":", 1)
                row_cells = table.add_row().cells
                row_cells[0].text = parts[0].strip()
                row_cells[1].text = parts[1].strip() if len(parts) > 1 else ""
            else:
                row_cells = table.add_row().cells
                row_cells[0].text = term
                row_cells[1].text = ""

        # Removed extra spacing at the end
